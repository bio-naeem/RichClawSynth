#!/usr/bin/env python3
"""
Generate synthetic test document files for benchmark queries.
Supports: CSV, XLSX (Excel), DOCX (Word).
"""

import argparse
import os
import csv
import random
import importlib
from datetime import datetime, timedelta

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SKILL_DIR = os.path.dirname(SCRIPT_DIR)
WORKSPACE_ROOT = os.path.dirname(os.path.dirname(SKILL_DIR))
DEFAULT_WORKSPACE = os.path.join(WORKSPACE_ROOT, 'tmp_output', 'claw-input-file-generator')
SETUP_COMMAND = 'pip3 install Pillow markdown weasyprint openpyxl python-docx gTTS'

# Sample data templates
PRODUCTS = ['笔记本电脑', '台式电脑', '显示器', '键盘', '鼠标', '耳机', '摄像头', 'U盘', '移动硬盘', '平板电脑']
REGIONS = ['华北', '华东', '华南', '华中', '西南', '西北', '东北']
NAMES = ['张三', '李四', '王五', '赵六', '钱七', '孙八', '周九', '吴十']
DEPARTMENTS = ['销售部', '技术部', '市场部', '人事部', '财务部']


def generate_csv(output_path, data_type='sales', rows=50):
    """Generate a CSV file with sample data."""
    ensure_parent_dir(output_path)
    
    if data_type == 'sales':
        headers = ['日期', '产品名称', '销售数量', '单价', '销售额', '区域', '销售人员']
        data = []
        base_date = datetime(2024, 1, 1)
        
        for i in range(rows):
            date = (base_date + timedelta(days=random.randint(0, 90))).strftime('%Y-%m-%d')
            product = random.choice(PRODUCTS)
            qty = random.randint(1, 100)
            price = round(random.uniform(100, 10000), 2)
            amount = round(qty * price, 2)
            region = random.choice(REGIONS)
            salesperson = random.choice(NAMES)
            data.append([date, product, qty, price, amount, region, salesperson])
    
    elif data_type == 'employees':
        headers = ['工号', '姓名', '部门', '入职日期', '基本工资', '绩效奖金', '总工资']
        data = []
        
        for i in range(rows):
            emp_id = f'EMP{1000 + i}'
            name = random.choice(NAMES)
            dept = random.choice(DEPARTMENTS)
            hire_date = (datetime(2020, 1, 1) + timedelta(days=random.randint(0, 1460))).strftime('%Y-%m-%d')
            base_salary = random.randint(8000, 30000)
            bonus = random.randint(0, 5000)
            total = base_salary + bonus
            data.append([emp_id, name, dept, hire_date, base_salary, bonus, total])
    
    else:
        headers = ['ID', '名称', '数值', '状态']
        data = [[i, f'项目{i}', random.randint(1, 100), random.choice(['进行中', '已完成', '待处理'])] for i in range(rows)]
    
    with open(output_path, 'w', newline='', encoding='utf-8-sig') as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerows(data)
    
    return output_path


def generate_xlsx(output_path, data_type='sales', rows=50, sheets=1):
    """Generate an Excel file with sample data."""
    require_python_package('openpyxl', 'openpyxl')
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter
    ensure_parent_dir(output_path)
    
    wb = Workbook()
    
    # Remove default sheet
    wb.remove(wb.active)
    
    # Styles
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='2C5AA0', end_color='2C5AA0', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    sheet_types = ['sales', 'employees', 'inventory'][:sheets]
    
    for dtype in sheet_types:
        ws = wb.create_sheet(title=dtype)
        
        if dtype == 'sales':
            headers = ['日期', '产品名称', '销售数量', '单价', '销售额', '区域', '销售人员']
            data = []
            base_date = datetime(2024, 1, 1)
            
            for i in range(rows):
                date = (base_date + timedelta(days=random.randint(0, 90))).strftime('%Y-%m-%d')
                product = random.choice(PRODUCTS)
                qty = random.randint(1, 100)
                price = round(random.uniform(100, 10000), 2)
                amount = round(qty * price, 2)
                region = random.choice(REGIONS)
                salesperson = random.choice(NAMES)
                data.append([date, product, qty, price, amount, region, salesperson])
        
        elif dtype == 'employees':
            headers = ['工号', '姓名', '部门', '入职日期', '基本工资', '绩效奖金', '总工资']
            data = []
            
            for i in range(rows):
                emp_id = f'EMP{1000 + i}'
                name = random.choice(NAMES)
                dept = random.choice(DEPARTMENTS)
                hire_date = (datetime(2020, 1, 1) + timedelta(days=random.randint(0, 1460))).strftime('%Y-%m-%d')
                base_salary = random.randint(8000, 30000)
                bonus = random.randint(0, 5000)
                total = base_salary + bonus
                data.append([emp_id, name, dept, hire_date, base_salary, bonus, total])
        
        else:  # inventory
            headers = ['商品编号', '商品名称', '库存数量', '单价', '库存金额', '状态']
            data = []
            
            for i in range(rows):
                item_id = f'ITEM{100 + i}'
                name = random.choice(PRODUCTS)
                qty = random.randint(0, 500)
                price = round(random.uniform(50, 5000), 2)
                value = round(qty * price, 2)
                status = '充足' if qty > 100 else '正常' if qty > 20 else '低库存' if qty > 0 else '缺货'
                data.append([item_id, name, qty, price, value, status])
        
        # Write headers
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        # Write data
        for row_idx, row_data in enumerate(data, 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Auto-adjust column widths
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 15
    
    wb.save(output_path)
    return output_path


def generate_docx(output_path, doc_type='report'):
    """Generate a Word document with sample content."""
    require_python_package('docx', 'python-docx')
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    ensure_parent_dir(output_path)
    
    doc = Document()
    
    if doc_type == 'report':
        # Title
        title = doc.add_heading('项目报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('2024年度工作总结报告')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Section 1
        doc.add_heading('一、项目概述', level=1)
        doc.add_paragraph(
            '本项目旨在提升公司内部管理效率，通过引入先进的数字化工具，'
            '优化业务流程，提高员工工作效率和满意度。'
        )
        
        # Section 2
        doc.add_heading('二、主要工作内容', level=1)
        
        doc.add_heading('2.1 系统开发', level=2)
        doc.add_paragraph(
            '完成了核心业务系统的开发工作，包括：'
        )
        doc.add_paragraph('• 用户管理模块', style='List Bullet')
        doc.add_paragraph('• 权限控制模块', style='List Bullet')
        doc.add_paragraph('• 数据分析模块', style='List Bullet')
        doc.add_paragraph('• 报表生成模块', style='List Bullet')
        
        doc.add_heading('2.2 测试与上线', level=2)
        doc.add_paragraph(
            '经过多轮测试，系统已顺利上线运行。测试覆盖率达到了95%以上，'
            '主要功能均通过了用户验收测试。'
        )
        
        # Section 3
        doc.add_heading('三、项目成果', level=1)
        
        # Add a table
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        headers = ['指标', '目标值', '实际值']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        data = [
            ['开发完成度', '100%', '100%'],
            ['测试覆盖率', '90%', '95%'],
            ['用户满意度', '80%', '88%'],
        ]
        for row_idx, row_data in enumerate(data, 1):
            for col_idx, value in enumerate(row_data):
                table.rows[row_idx].cells[col_idx].text = value
        
        doc.add_paragraph()  # Spacer
        
        # Section 4
        doc.add_heading('四、总结与展望', level=1)
        doc.add_paragraph(
            '本项目圆满完成了预期目标，系统运行稳定，用户反馈良好。'
            '未来将继续优化系统功能，提升用户体验。'
        )
        
    elif doc_type == 'contract':
        # Title
        title = doc.add_heading('服务合同', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Parties
        doc.add_paragraph('甲方：XX科技有限公司')
        doc.add_paragraph('乙方：XX服务有限公司')
        doc.add_paragraph()
        
        # Content
        doc.add_heading('第一条 服务内容', level=1)
        doc.add_paragraph(
            '乙方向甲方提供以下服务：技术咨询、系统维护、培训服务等。'
        )
        
        doc.add_heading('第二条 服务期限', level=1)
        doc.add_paragraph('本合同有效期为一年，自2024年1月1日起至2024年12月31日止。')
        
        doc.add_heading('第三条 服务费用', level=1)
        doc.add_paragraph('服务费用总计人民币壹拾万元整（¥100,000.00）。')
        
        doc.add_heading('第四条 违约责任', level=1)
        doc.add_paragraph('任何一方违反本合同约定，应承担相应的违约责任。')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Signatures
        doc.add_paragraph('甲方（盖章）：________________')
        doc.add_paragraph('日期：________________')
        doc.add_paragraph()
        doc.add_paragraph('乙方（盖章）：________________')
        doc.add_paragraph('日期：________________')
        
    else:  # letter
        title = doc.add_heading('通知函', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph('致：全体员工')
        doc.add_paragraph()
        
        doc.add_paragraph('主题：关于调整工作时间安排的通知')
        doc.add_paragraph()
        
        doc.add_paragraph('各位同事：')
        doc.add_paragraph()
        doc.add_paragraph(
            '根据公司发展需要，经管理层研究决定，自2024年7月1日起，'
            '公司工作时间调整如下：'
        )
        doc.add_paragraph()
        doc.add_paragraph('• 上午工作时间：9:00 - 12:00')
        doc.add_paragraph('• 下午工作时间：13:30 - 18:00')
        doc.add_paragraph()
        doc.add_paragraph('请各位同事做好相应安排，特此通知。')
        doc.add_paragraph()
        doc.add_paragraph('XX科技有限公司')
        doc.add_paragraph('2024年6月15日')
    
    doc.save(output_path)
    return output_path


def ensure_parent_dir(path):
    parent = os.path.dirname(os.path.abspath(path))
    if parent:
        os.makedirs(parent, exist_ok=True)


def require_python_package(module_name, package_name):
    try:
        importlib.import_module(module_name)
    except ImportError as exc:
        raise SystemExit(
            f"Missing Python dependency '{package_name}'. "
            f"Install skill prerequisites first: {SETUP_COMMAND}"
        ) from exc


def main():
    parser = argparse.ArgumentParser(description='Generate synthetic test document files')
    parser.add_argument('--type', choices=['csv', 'xlsx', 'docx'], required=True,
                       help='Document type to generate')
    parser.add_argument('--output', default=None, help='Output file path')
    parser.add_argument('--workspace', default=DEFAULT_WORKSPACE, help='Workspace directory')
    parser.add_argument('--data-type', choices=['sales', 'employees', 'inventory'],
                       default='sales', help='Type of sample data')
    parser.add_argument('--doc-type', choices=['report', 'contract', 'letter'],
                       default='report', help='Document template type')
    parser.add_argument('--rows', type=int, default=50, help='Number of data rows')
    parser.add_argument('--sheets', type=int, default=1, help='Number of sheets (xlsx)')
    
    args = parser.parse_args()
    
    # Determine output path
    if args.output:
        output_path = args.output
    else:
        os.makedirs(args.workspace, exist_ok=True)
        ext = {'csv': '.csv', 'xlsx': '.xlsx', 'docx': '.docx'}[args.type]
        filename = f'{args.data_type}_data{ext}' if args.type != 'docx' else f'{args.doc_type}{ext}'
        output_path = os.path.join(args.workspace, filename)
    
    # Generate file
    if args.type == 'csv':
        generate_csv(output_path, args.data_type, args.rows)
    elif args.type == 'xlsx':
        generate_xlsx(output_path, args.data_type, args.rows, args.sheets)
    elif args.type == 'docx':
        generate_docx(output_path, args.doc_type)

    if not os.path.exists(output_path):
        raise SystemExit(f"Generator did not create the expected file: {output_path}")

    print(f"Generated: {output_path}")


if __name__ == '__main__':
    main()
